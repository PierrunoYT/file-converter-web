from flask import Flask, request, send_file, abort
from PIL import Image
import io
import os
import logging
from flask_cors import CORS
import pillow_heif
import pillow_avif  # Import to register AVIF plugin
from pydub import AudioSegment
from moviepy.video.io.VideoFileClip import VideoFileClip
import tempfile
from werkzeug.utils import secure_filename
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import docx
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
import odf
from odf.opendocument import OpenDocumentText
from odf import text as odf_text
import markdown
from docx2pdf import convert

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100MB max file size
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp', 'heic', 'avif'}
ALLOWED_AUDIO_EXTENSIONS = {'mp3', 'wav', 'ogg', 'flac', 'm4a', 'aac'}
ALLOWED_VIDEO_EXTENSIONS = {'mp4', 'avi', 'mov', 'mkv', 'webm'}
ALLOWED_TEXT_EXTENSIONS = {'txt', 'doc', 'docx', 'pdf', 'rtf', 'odt', 'md'}

# Register HEIF opener with Pillow
pillow_heif.register_heif_opener()

app = Flask(__name__, static_url_path='', static_folder='static')
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
CORS(app)  # Enable CORS for all routes

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def validate_file(file, allowed_extensions):
    if not file:
        abort(400, description="No file provided")
    filename = secure_filename(file.filename)
    if not allowed_file(filename, allowed_extensions):
        abort(400, description=f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}")
    return filename

# Serve static files
@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/image-converter')
def image_converter():
    return app.send_static_file('image-converter.html')

@app.route('/audio-converter')
def audio_converter():
    return app.send_static_file('audio-converter.html')

@app.route('/video-converter')
def video_converter():
    return app.send_static_file('video-converter.html')

@app.route('/text-converter')
def text_converter():
    return app.send_static_file('text-converter.html')

# Serve other static files
@app.route('/<path:path>')
def serve_static(path):
    return app.send_static_file(path)

@app.route('/convert-video', methods=['POST'])
@limiter.limit("10 per minute")
def convert_video():
    temp_files = []  # Keep track of temporary files
    try:
        # Get and validate the video file
        video_file = request.files.get('file')
        filename = validate_file(video_file, ALLOWED_VIDEO_EXTENSIONS)

        target_format = request.form.get('targetFormat')
        if not target_format or target_format not in ALLOWED_VIDEO_EXTENSIONS:
            abort(400, description=f"Invalid target format. Allowed formats: {', '.join(ALLOWED_VIDEO_EXTENSIONS)}")

        quality = request.form.get('quality', 'high')  # high, medium, low

        logger.info(f"Starting video conversion: {filename} to {target_format}")

        # Create temporary files with unique names
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{video_file.filename.split(".")[-1]}')
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{target_format}')
        temp_files.extend([temp_input.name, temp_output.name])

        # Immediately close the file handles
        temp_input.close()
        temp_output.close()

        # Save the uploaded video
        video_file.save(temp_input.name)

        try:
            # Load the video file
            video = VideoFileClip(temp_input.name)

            # Set bitrate based on quality
            bitrate = {
                'high': '8000k',
                'medium': '4000k',
                'low': '2000k'
            }.get(quality, '4000k')

            # Write the converted video
            video.write_videofile(
                temp_output.name,
                codec='libx264' if target_format == 'mp4' else None,
                bitrate=bitrate,
                audio_codec='aac' if target_format in ['mp4', 'mov'] else 'libvorbis',
                preset='medium',
                threads=4
            )

            # Close the video to free up resources
            video.close()

            # Send the converted video file
            return send_file(
                temp_output.name,
                mimetype=f'video/{target_format}',
                as_attachment=True,
                download_name=f'converted_video.{target_format}'
            )

        finally:
            # Clean up temporary files
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.close(os.open(temp_file, os.O_RDONLY))  # Ensure file handle is closed
                        os.unlink(temp_file)
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file {temp_file}: {str(e)}")

    except Exception as e:
        print(f"Error converting video: {str(e)}")
        return str(e), 500

@app.route('/convert-audio', methods=['POST'])
@limiter.limit("20 per minute")
def convert_audio():
    try:
        # Get and validate the audio file
        audio_file = request.files.get('audio')
        filename = validate_file(audio_file, ALLOWED_AUDIO_EXTENSIONS)

        target_format = request.form.get('format')
        if not target_format or target_format.lower() not in ALLOWED_AUDIO_EXTENSIONS:
            abort(400, description=f"Invalid target format. Allowed formats: {', '.join(ALLOWED_AUDIO_EXTENSIONS)}")

        original_filename = secure_filename(request.form.get('filename', 'converted_audio'))

        logger.info(f"Starting audio conversion: {filename} to {target_format}")

        try:
            # Read the audio file using pydub
            audio = AudioSegment.from_file(audio_file)

            # Prepare output buffer
            output_buffer = io.BytesIO()

            # Export the audio to the target format
            audio.export(output_buffer, format=target_format.lower())
            output_buffer.seek(0)

            # Send the converted audio
            return send_file(
                output_buffer,
                mimetype=f'audio/{target_format.lower()}',
                as_attachment=True,
                download_name=f'{original_filename}.{target_format.lower()}'
            )

        except Exception as e:
            error_msg = f"Error during audio conversion: {str(e)}"
            logger.error(error_msg)
            return {"error": error_msg}, 500

    except Exception as e:
        error_msg = f"Error processing audio request: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}, 500

@app.route('/convert', methods=['POST'])
@limiter.limit("30 per minute")
def convert_image():
    try:
        # Get and validate the image file
        image_file = request.files.get('image')
        filename = validate_file(image_file, ALLOWED_IMAGE_EXTENSIONS)
        
        target_format = request.form.get('format')
        if not target_format or target_format.lower() not in ALLOWED_IMAGE_EXTENSIONS:
            abort(400, description=f"Invalid target format. Allowed formats: {', '.join(ALLOWED_IMAGE_EXTENSIONS)}")
            
        compression = request.form.get('compression', '95')  # New compression parameter
        try:
            compression = int(compression)
            if not (0 <= compression <= 100):
                raise ValueError
        except ValueError:
            compression = 95
            
        original_filename = secure_filename(request.form.get('filename', 'converted_image'))
        
        logger.info(f"Starting image conversion: {filename} to {target_format}")

        # Read the image using PIL
        image = Image.open(image_file)

        # Convert RGBA to RGB if needed
        if image.mode == 'RGBA' and target_format.lower() in ['jpg', 'heic']:
            background = Image.new('RGB', image.size, (255, 255, 255))
            background.paste(image, mask=image.split()[3])
            image = background

        # Prepare output buffer
        output_buffer = io.BytesIO()

        # Save the converted image to the buffer
        if target_format.lower() == 'jpg':
            image.save(output_buffer, 'JPEG', quality=95)
        elif target_format.lower() == 'heic':
            # Convert to RGB if not already
            if image.mode != 'RGB':
                image = image.convert('RGB')
            # Save as HEIC
            pillow_heif.save(output_buffer, image, quality=95)
        elif target_format.lower() == 'avif':
            # Convert to RGB if not already
            if image.mode != 'RGB':
                image = image.convert('RGB')
            image.save(output_buffer, 'AVIF', quality=95, speed=6)
        elif target_format.lower() == 'tiff':
            image.save(output_buffer, 'TIFF', compression='tiff_lzw')
        elif target_format.lower() == 'bmp':
            image.save(output_buffer, 'BMP')
        else:
            image.save(output_buffer, target_format.upper())

        output_buffer.seek(0)

        # Send the converted image
        return send_file(
            output_buffer,
            mimetype=f'image/{target_format.lower()}',
            as_attachment=True,
            download_name=f'{original_filename}.{target_format.lower()}'
        )

    except Exception as e:
        error_msg = f"Error converting image: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}, 500

@app.route('/convert/text', methods=['POST'])
@limiter.limit("20 per minute")
def convert_text():
    temp_files = []  # Keep track of temporary files
    try:
        # Get and validate the text file
        text_file = request.files.get('file')
        filename = validate_file(text_file, ALLOWED_TEXT_EXTENSIONS)

        target_format = request.form.get('target_format')
        if not target_format or target_format.lower() not in ALLOWED_TEXT_EXTENSIONS:
            abort(400, description=f"Invalid target format. Allowed formats: {', '.join(ALLOWED_TEXT_EXTENSIONS)}")

        # Create temporary directory for all temporary files
        temp_dir = tempfile.mkdtemp()
        temp_input_path = os.path.join(temp_dir, f'input.{filename.split(".")[-1]}')
        temp_output_path = os.path.join(temp_dir, f'output.{target_format}')
        temp_files.extend([temp_input_path, temp_output_path])
        
        try:
            # Save input file
            text_file.save(temp_input_path)

            # Convert based on input and target formats
            if target_format == 'txt':
                # Convert to plain text
                if filename.endswith('.md'):
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()
                    # Convert markdown to plain text (strips all markdown formatting)
                    html_content = markdown.markdown(md_content)
                    # Simple HTML to text conversion (strips HTML tags)
                    text_content = html_content.replace('<p>', '').replace('</p>', '\n\n')
                    text_content = text_content.replace('<h1>', '').replace('</h1>', '\n\n')
                    text_content = text_content.replace('<h2>', '').replace('</h2>', '\n\n')
                    text_content = text_content.replace('<h3>', '').replace('</h3>', '\n\n')
                    text_content = text_content.replace('<ul>', '\n').replace('</ul>', '\n')
                    text_content = text_content.replace('<li>', '* ').replace('</li>', '\n')
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                elif filename.endswith('.docx'):
                    doc = Document(temp_input_path)
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        for para in doc.paragraphs:
                            f.write(para.text + '\n')
                elif filename.endswith('.pdf'):
                    reader = PdfReader(temp_input_path)
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        for page in reader.pages:
                            f.write(page.extract_text() + '\n')
                elif filename.endswith('.odt'):
                    doc = OpenDocumentText(temp_input_path)
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        for element in doc.getElementsByType(odf_text.P):
                            f.write(element.text + '\n')
                else:
                    # For txt files, just copy
                    with open(temp_input_path, 'r', encoding='utf-8') as f_in:
                        with open(temp_output_path, 'w', encoding='utf-8') as f_out:
                            f_out.write(f_in.read())
            
            elif target_format == 'docx':
                # Convert to DOCX
                doc = Document()
                if filename.endswith('.md'):
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()
                    # Convert markdown to HTML
                    html_content = markdown.markdown(md_content)
                    # Parse the HTML content and add to document with formatting
                    # This is a simple conversion that maintains basic formatting
                    for line in html_content.split('\n'):
                        if line.startswith('<h1>'):
                            doc.add_heading(line[4:-5], level=1)
                        elif line.startswith('<h2>'):
                            doc.add_heading(line[4:-5], level=2)
                        elif line.startswith('<h3>'):
                            doc.add_heading(line[4:-5], level=3)
                        elif line.startswith('<p>'):
                            doc.add_paragraph(line[3:-4])
                        elif line.startswith('<li>'):
                            doc.add_paragraph(line[4:-5], style='List Bullet')
                elif filename.endswith('.txt'):
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            doc.add_paragraph(line.strip())
                elif filename.endswith('.pdf'):
                    reader = PdfReader(temp_input_path)
                    for page in reader.pages:
                        doc.add_paragraph(page.extract_text())
                doc.save(temp_output_path)
            
            elif target_format == 'pdf':
                # Convert to PDF
                writer = PdfWriter()
                if filename.endswith('.md'):
                    # Convert markdown to docx first
                    doc = Document()
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()
                    # Convert markdown to HTML
                    html_content = markdown.markdown(md_content)
                    # Parse the HTML content and add to document with formatting
                    for line in html_content.split('\n'):
                        if line.startswith('<h1>'):
                            doc.add_heading(line[4:-5], level=1)
                        elif line.startswith('<h2>'):
                            doc.add_heading(line[4:-5], level=2)
                        elif line.startswith('<h3>'):
                            doc.add_heading(line[4:-5], level=3)
                        elif line.startswith('<p>'):
                            doc.add_paragraph(line[3:-4])
                        elif line.startswith('<li>'):
                            doc.add_paragraph(line[4:-5], style='List Bullet')
                    temp_docx_path = os.path.join(temp_dir, 'temp.docx')
                    doc.save(temp_docx_path)
                    # Convert DOCX to PDF
                    convert(temp_docx_path, temp_output_path)
                    temp_files.append(temp_docx_path)  # Add to cleanup list
                elif filename.endswith('.txt'):
                    # Create PDF from text
                    doc = Document()
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            doc.add_paragraph(line.strip())
                    temp_docx_path = os.path.join(temp_dir, 'temp.docx')
                    doc.save(temp_docx_path)
                    # Convert DOCX to PDF
                    convert(temp_docx_path, temp_output_path)
                    temp_files.append(temp_docx_path)  # Add to cleanup list

            elif target_format == 'md':
                # Convert to Markdown
                if filename.endswith('.txt'):
                    # Simple conversion: each line becomes a paragraph
                    with open(temp_input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    # Convert plain text to markdown (each paragraph gets wrapped in markdown)
                    md_content = '\n\n'.join(f'{para}' for para in content.split('\n\n'))
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
                elif filename.endswith('.docx'):
                    doc = Document(temp_input_path)
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        for para in doc.paragraphs:
                            # Convert heading styles to markdown
                            if para.style.name.startswith('Heading'):
                                level = int(para.style.name[-1])
                                f.write('#' * level + ' ' + para.text + '\n\n')
                            # Convert list items
                            elif para.style.name == 'List Bullet':
                                f.write('* ' + para.text + '\n')
                            # Regular paragraphs
                            else:
                                f.write(para.text + '\n\n')
                elif filename.endswith('.pdf'):
                    reader = PdfReader(temp_input_path)
                    with open(temp_output_path, 'w', encoding='utf-8') as f:
                        for page in reader.pages:
                            # Convert PDF text to markdown paragraphs
                            text = page.extract_text()
                            paragraphs = text.split('\n\n')
                            for para in paragraphs:
                                f.write(para.strip() + '\n\n')

            # Create a copy of the output file in memory
            with open(temp_output_path, 'rb') as f:
                file_content = io.BytesIO(f.read())
                file_content.seek(0)

            # Close any open file handles
            if 'doc' in locals():
                del doc
            if 'reader' in locals():
                del reader

            # Send the file from memory
            return send_file(
                file_content,
                mimetype=f'application/{target_format}',
                as_attachment=True,
                download_name=f'converted_document.{target_format}'
            )

        finally:
            # Clean up files
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        # Try to close any remaining handles on Windows
                        try:
                            os.close(os.open(temp_file, os.O_RDONLY))
                        except:
                            pass
                        os.unlink(temp_file)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_file}: {str(e)}")

            # Clean up temporary directory
            try:
                if os.path.exists(temp_dir):
                    import time
                    time.sleep(0.1)  # Small delay to allow Windows to release handles
                    # Try to remove any remaining files
                    for root, dirs, files in os.walk(temp_dir):
                        for name in files:
                            try:
                                os.unlink(os.path.join(root, name))
                            except:
                                pass
                    os.rmdir(temp_dir)
            except Exception as e:
                logger.warning(f"Failed to delete temporary directory {temp_dir}: {str(e)}")

    except Exception as e:
        # Cleanup on error
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    # Try to close any remaining handles on Windows
                    try:
                        os.close(os.open(temp_file, os.O_RDONLY))
                    except:
                        pass
                    os.unlink(temp_file)
            except:
                pass
        try:
            if os.path.exists(temp_dir):
                import time
                time.sleep(0.1)  # Small delay to allow Windows to release handles
                # Try to remove any remaining files
                for root, dirs, files in os.walk(temp_dir):
                    for name in files:
                        try:
                            os.unlink(os.path.join(root, name))
                        except:
                            pass
                os.rmdir(temp_dir)
        except:
            pass

        error_msg = f"Error converting text document: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}, 500

if __name__ == '__main__':
    # Ensure the static folder exists
    os.makedirs('static', exist_ok=True)
    
    # Copy frontend files to static folder if they exist
    import shutil
    for file in ['index.html', 'style.css', 'script.js']:
        if os.path.exists(file):
            shutil.copy(file, os.path.join('static', file))
    
    app.run(debug=True, port=5000)