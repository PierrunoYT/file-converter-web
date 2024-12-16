from flask import request, send_file, abort
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from odf.opendocument import OpenDocumentText
from odf import text as odf_text
import markdown
from docx2pdf import convert
import tempfile
import os
import io
from utils import validate_file, cleanup_temp_files, cleanup_temp_dir
from config import ALLOWED_TEXT_EXTENSIONS, logger

def handle_text_conversion():
    temp_files = []  # Keep track of temporary files
    temp_dir = None
    try:
        # Get and validate the text file
        text_file = request.files.get('file')
        filename = validate_file(text_file, ALLOWED_TEXT_EXTENSIONS)

        target_format = request.form.get('target_format')
        if not target_format or target_format.lower() not in ALLOWED_TEXT_EXTENSIONS:
            abort(400, description=f"Invalid target format. Allowed formats: {', '.join(ALLOWED_TEXT_EXTENSIONS)}")

        # Create temporary directory for all temporary files
        temp_dir = tempfile.mkdtemp()
        temp_input_path = os.path.join(temp_dir, f'input.{filename.split(".")[-1]}')
        temp_output_path = os.path.join(temp_dir, f'output.{target_format}')
        temp_files.extend([temp_input_path, temp_output_path])
        
        try:
            # Save input file
            text_file.save(temp_input_path)

            # Convert based on input and target formats
            if target_format == 'txt':
                convert_to_txt(filename, temp_input_path, temp_output_path)
            elif target_format == 'docx':
                convert_to_docx(filename, temp_input_path, temp_output_path)
            elif target_format == 'pdf':
                convert_to_pdf(filename, temp_input_path, temp_output_path, temp_dir, temp_files)
            elif target_format == 'md':
                convert_to_md(filename, temp_input_path, temp_output_path)

            # Create a copy of the output file in memory
            with open(temp_output_path, 'rb') as f:
                file_content = io.BytesIO(f.read())
                file_content.seek(0)

            # Send the file from memory
            return send_file(
                file_content,
                mimetype=f'application/{target_format}',
                as_attachment=True,
                download_name=f'converted_document.{target_format}'
            )

        finally:
            # Clean up files and directory
            cleanup_temp_files(temp_files)
            if temp_dir:
                cleanup_temp_dir(temp_dir)

    except Exception as e:
        # Cleanup on error
        if temp_files:
            cleanup_temp_files(temp_files)
        if temp_dir:
            cleanup_temp_dir(temp_dir)

        error_msg = f"Error converting text document: {str(e)}"
        logger.error(error_msg)
        return {"error": error_msg}, 500

def convert_to_txt(filename, input_path, output_path):
    if filename.endswith('.md'):
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        # Convert markdown to plain text (strips all markdown formatting)
        html_content = markdown.markdown(md_content)
        # Simple HTML to text conversion (strips HTML tags)
        text_content = html_content.replace('<p>', '').replace('</p>', '\n\n')
        text_content = text_content.replace('<h1>', '').replace('</h1>', '\n\n')
        text_content = text_content.replace('<h2>', '').replace('</h2>', '\n\n')
        text_content = text_content.replace('<h3>', '').replace('</h3>', '\n\n')
        text_content = text_content.replace('<ul>', '\n').replace('</ul>', '\n')
        text_content = text_content.replace('<li>', '* ').replace('</li>', '\n')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
    elif filename.endswith('.docx'):
        doc = Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
    elif filename.endswith('.pdf'):
        reader = PdfReader(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for page in reader.pages:
                f.write(page.extract_text() + '\n')
    elif filename.endswith('.odt'):
        doc = OpenDocumentText(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for element in doc.getElementsByType(odf_text.P):
                f.write(element.text + '\n')
    else:
        # For txt files, just copy
        with open(input_path, 'r', encoding='utf-8') as f_in:
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(f_in.read())

def convert_to_docx(filename, input_path, output_path):
    doc = Document()
    if filename.endswith('.md'):
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content)
        # Parse the HTML content and add to document with formatting
        for line in html_content.split('\n'):
            if line.startswith('<h1>'):
                doc.add_heading(line[4:-5], level=1)
            elif line.startswith('<h2>'):
                doc.add_heading(line[4:-5], level=2)
            elif line.startswith('<h3>'):
                doc.add_heading(line[4:-5], level=3)
            elif line.startswith('<p>'):
                doc.add_paragraph(line[3:-4])
            elif line.startswith('<li>'):
                doc.add_paragraph(line[4:-5], style='List Bullet')
    elif filename.endswith('.txt'):
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.strip())
    elif filename.endswith('.pdf'):
        reader = PdfReader(input_path)
        for page in reader.pages:
            doc.add_paragraph(page.extract_text())
    doc.save(output_path)

def convert_to_pdf(filename, input_path, output_path, temp_dir, temp_files):
    if filename.endswith('.md') or filename.endswith('.txt'):
        # Convert to docx first
        doc = Document()
        if filename.endswith('.md'):
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            # Convert markdown to HTML
            html_content = markdown.markdown(md_content)
            # Parse the HTML content and add to document with formatting
            for line in html_content.split('\n'):
                if line.startswith('<h1>'):
                    doc.add_heading(line[4:-5], level=1)
                elif line.startswith('<h2>'):
                    doc.add_heading(line[4:-5], level=2)
                elif line.startswith('<h3>'):
                    doc.add_heading(line[4:-5], level=3)
                elif line.startswith('<p>'):
                    doc.add_paragraph(line[3:-4])
                elif line.startswith('<li>'):
                    doc.add_paragraph(line[4:-5], style='List Bullet')
        else:  # txt file
            with open(input_path, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.strip())
        
        temp_docx_path = os.path.join(temp_dir, 'temp.docx')
        doc.save(temp_docx_path)
        temp_files.append(temp_docx_path)
        # Convert DOCX to PDF
        convert(temp_docx_path, output_path)

def convert_to_md(filename, input_path, output_path):
    if filename.endswith('.txt'):
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        # Convert plain text to markdown (each paragraph gets wrapped in markdown)
        md_content = '\n\n'.join(f'{para}' for para in content.split('\n\n'))
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
    elif filename.endswith('.docx'):
        doc = Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                # Convert heading styles to markdown
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    f.write('#' * level + ' ' + para.text + '\n\n')
                # Convert list items
                elif para.style.name == 'List Bullet':
                    f.write('* ' + para.text + '\n')
                # Regular paragraphs
                else:
                    f.write(para.text + '\n\n')
    elif filename.endswith('.pdf'):
        reader = PdfReader(input_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for page in reader.pages:
                # Convert PDF text to markdown paragraphs
                text = page.extract_text()
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    f.write(para.strip() + '\n\n')